from rest_framework.decorators import api_view
from rest_framework.response import Response
from rest_framework import status
from .models import Report , DocumentFormat, Paragraph, Text, Section,Image,ProductReport,ShopDrawReport
from .serializers import ReportSerializer , DocumentFormatSerializer , ProductReportSerializer ,SectionSerializer,ImageSerializer ,ShopDrawReportSerializer
from django.contrib.auth.models import User , Group
from rest_framework.authtoken.models import Token
from django.http import HttpResponse 
from django.core.serializers import serialize
from docx.shared import Pt ,Inches
from docx import Document
import json
from io import BytesIO
import pandas as pd
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import zipfile
from django.views.decorators.csrf import csrf_exempt
from corsheaders.defaults import default_headers
from django.shortcuts import render
from django.conf import settings
from docx2pdf import convert
import tempfile
import os
import pythoncom


SL_TYPEMAP = {"ผลิตภัณฑ์": ProductReportSerializer , "SHOP DRAWING": ShopDrawReportSerializer}
TYPEMAP = {"ผลิตภัณฑ์": ProductReport , "SHOP DRAWING": ShopDrawReport}

def index(request):
    # Directly render the index.html from the build directory
    return render(request, os.path.join(settings.BASE_DIR, 'static', 'react', 'dist', 'index.html'))

@api_view(['POST'])
@csrf_exempt
def register_user(request):
    username = request.data.get('username')
    password = request.data.get('password')
    if User.objects.filter(username=username).exists():
        return Response({'error': 'Username already taken'}, status=status.HTTP_400_BAD_REQUEST)
        # Create the user
    user = User.objects.create_user(username=username, password=password)
    user.groups.add(Group.objects.get(name='Pre_Register'))
    return Response(status=status.HTTP_201_CREATED)

@api_view(['POST'])
@csrf_exempt
def login(request):
    username = request.data.get('username')
    password = request.data.get('password')  
    if User.objects.filter(username=username).exists():
        user = User.objects.get(username = username)
        if user.check_password(password):  # Use `check_password` for hashed passwords
            token, _ = Token.objects.get_or_create(user=user)
            return Response({"message": "Login successful", "token": token.key,'role': [group.name for group in user.groups.all()]}, status=200)
        else:
            return Response({"error": "Invalid password."}, status=401)
    else:
        return Response({"error": "User does not exist."}, status=404)

def gen_download_report(json_data):
    doc1 , doc2 = get_doc(json_data)
    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        doc1_buffer = BytesIO()
        doc1.save(doc1_buffer)
        doc1_buffer.seek(0)
        zip_file.writestr('บันทึกข้อความ.docx', doc1_buffer.getvalue())    
        '''
        pdf1_buffer = convert_docx_to_pdf(doc1_buffer)
        zip_file.writestr('บันทึกข้อความ.pdf', pdf1_buffer.getvalue())  
        ''' 
        doc2_buffer = BytesIO()
        doc2.save(doc2_buffer)
        doc2_buffer.seek(0)
        zip_file.writestr('ผลการพิจารณา.doc', doc2_buffer.getvalue())
        '''
        pdf2_buffer = convert_docx_to_pdf(doc2_buffer)
        zip_file.writestr('ผลการพิจารณา.pdf', pdf2_buffer.getvalue()) 
        '''
        
        zip_buffer.seek(0)
    return zip_buffer
    
@api_view(['POST'])
@csrf_exempt
def create_report(request):
    json_data = request.data['report']
    isDownload = request.data['download']
    token = Token.objects.get(key=request.data['userToken'])
    type = json_data['type']
    serializer = SL_TYPEMAP[type](data=json_data)
    if serializer.is_valid():
        serializer.save(create_by = token.user)
        if(isDownload):  
            zip_buffer = gen_download_report(json_data)
            response = HttpResponse(
                zip_buffer,
                content_type='application/zip'
            )
            response['Content-Disposition'] = 'attachment; filename="report.zip"'
            return response
           
        else:
            return Response(serializer.data, status=status.HTTP_200_OK)
    else:
        print(serializer.error_messages)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    
@api_view(['GET'])
@csrf_exempt
def get_all_report(request):
    token = request.query_params.get('token')
    report = Report.objects.all()
    serializer = ReportSerializer(report, many=True)
    data_to_send = serializer.data
    creater = Token.objects.get(key=token).user.username
    for report in data_to_send:
        if('create_by' in report.keys()):
            if(report['create_by'] == creater):
                report['isCreate'] = True
            else:
                report['isCreate'] = False
    return Response(data_to_send, status=status.HTTP_200_OK)

@api_view(['GET'])
@csrf_exempt
def get_select_report(request):
    type = request.query_params.get('type')
    id = request.query_params.get('id')
    report = TYPEMAP[type].objects.get(id = id)
    serializer = SL_TYPEMAP[type](report)
    return Response(serializer.data, status=status.HTTP_200_OK)

@api_view(['GET'])
@csrf_exempt
def dowload_select_report(request):
    try:
        type = request.query_params.get('type')
        id = request.query_params.get('id')
        report = TYPEMAP[type].objects.get(id = id)
        serializer = SL_TYPEMAP[type](report)
        zip_buffer = gen_download_report(serializer.data)
        response = HttpResponse(
            zip_buffer,
            content_type='application/zip'
        )
        response['Content-Disposition'] = 'attachment; filename="report.zip"'
        return response
    except Exception as e:
        print(e)
        return Response(status=status.HTTP_400_BAD_REQUEST)
    
@api_view(['POST'])
@csrf_exempt
def update_select_report(request):
    try:
        
        json_data = request.data['report']
        isDownload = request.data['download']
        type = request.data['type']
        id = request.data['id']
        report = TYPEMAP[type].objects.get(id=id)  
        for key, value in json_data.items():
            setattr(report, key, value)

# Save the changes to the database
        report.save()
        serializer = SL_TYPEMAP[type](report)
        if(isDownload):
            zip_buffer = BytesIO()
            zip_buffer = gen_download_report(serializer.data)
            response = HttpResponse(
                zip_buffer,
                content_type='application/zip'
            )
            response['Content-Disposition'] = 'attachment; filename="report.zip"'
            return response
        else:
            return Response(serializer.data, status=status.HTTP_200_OK)
    except Exception as e:
        print(e)
        return Response(status=status.HTTP_400_BAD_REQUEST)

@api_view(['DELETE'])
@csrf_exempt
def delete_report(request):
    try:
        id = request.data
        Report.objects.get(id =id).delete()
        return Response(status=status.HTTP_204_NO_CONTENT)
    except Report.DoesNotExist:
        return Response(status=status.HTTP_404_NOT_FOUND)
    except Exception as e:
        return Response(status=status.HTTP_400_BAD_REQUEST)

@api_view(['GET'])
@csrf_exempt
def get_all_user(request):
    try:
        users = User.objects.all()
        users_json = serialize('json', users)
        return Response(users_json, status=status.HTTP_200_OK)
    except Exception as e:
        return Response(status=status.HTTP_400_BAD_REQUEST)

@api_view(['DELETE'])
@csrf_exempt
def delete_user(request):
    try:
        username = request.data
        User.objects.get(username = username).delete()
        return Response(status=status.HTTP_204_NO_CONTENT)
    except Report.DoesNotExist:
        return Response(status=status.HTTP_404_NOT_FOUND)
    except Exception as e:
        return Response(status=status.HTTP_400_BAD_REQUEST)

@api_view(['POST'])
@csrf_exempt
def user_update_role(request):
    try:
        json_data= request.data
        user  = User.objects.get(username = json_data['username'])
        user.groups.clear()
        user.groups.add(Group.objects.get(name=json_data['role']))
        return Response(status=status.HTTP_200_OK)
    except Report.DoesNotExist:
        return Response(status=status.HTTP_404_NOT_FOUND)
    except Exception as e:
        return Response(status=status.HTTP_400_BAD_REQUEST)
    
@api_view(['GET'])
@csrf_exempt
def get_all_format(request):
    try:
        formats = DocumentFormat.objects.all()
        formats_json = serialize('json', formats )
        return Response(formats_json , status=status.HTTP_200_OK)
    except Exception as e:
        print(e)
        return Response(status=status.HTTP_400_BAD_REQUEST)
    
@api_view(['POST'])
@csrf_exempt
def update_format(request):
    if request.method == 'POST' and request.FILES.get('file'):
        try:
            name = request.POST.get("itemName")
            DocumentFormat.objects.filter(name = name).delete()
            return add_form(request._request)
        except Report.DoesNotExist:
            return Response(status=status.HTTP_404_NOT_FOUND)
        except Exception as e:
            return Response(status=status.HTTP_400_BAD_REQUEST)
    
@api_view(['GET'])
@csrf_exempt
def dowload_select_format(request):
    try:
        name = request.query_params.get('name')
        select  = DocumentFormat.objects.filter(name = name).first()

        if select.name == 'product':
            doc = print_product_report()
        elif select.name == 'ShopDrawResult':
           doc = print_shop_draw_result()
        else:
            doc = Document()
            create_doc_section(doc, Section.objects.first(),False)
            if select.name == 'product_result_head':
                head = DocumentFormat.objects.filter(name = 'product_result_head').first()
                for paragraph_data in head.paragraphs.all():
                    create_doc_paragraph(doc, paragraph_data,True)
            if select.name == 'product_result_fristPara':
                fristPara = DocumentFormat.objects.filter(name = 'product_result_fristPara').first()
                for paragraph_data in fristPara.paragraphs.all():
                    create_doc_paragraph(doc, paragraph_data,True)
            if select.name == 'product_result_tail':
                tail = DocumentFormat.objects.filter(name = 'product_result_tail').first()
                for paragraph_data in tail.paragraphs.all():
                    create_doc_paragraph(doc, paragraph_data,True)
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # Create a response with the buffer content
        response = HttpResponse(
            buffer,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename=".docx"'
        return response
    except Exception as e:
        print(e)
        return Response(status=status.HTTP_400_BAD_REQUEST)
    
      
      
    
'''
@api_view(['GET'])
@csrf_exempt
def check_miss_spell(request):
    try:
        text= request.query_params.get('text')
        words = word_tokenize(text, engine="wordcut")
        missSpells = []
        count = 1
        for word in words:
            suggestions = spell(word)
            print(count , word)
            count = count +1
            if suggestions and word not in suggestions:
                missSpells.append( {"word": word,"suggestions": suggestions})
        return JsonResponse({
            "success": True,
            "misspellings": missSpells,
            "total_words": len(words),
            "total_errors": len(missSpells)
        })
    except Exception as e:
        
        return Response(status=status.HTTP_400_BAD_REQUEST);
'''


#Behide is dev tool




@api_view(['POST'])
def add_section(request):
    if request.method == 'POST' and request.FILES.get('file'):
        file = request.FILES['file']
        try:
            # Read the content of the JSON file
            doc = Document(file)
            sec = doc.sections[0]
            newSec = Section.objects.create(
                top = sec.top_margin,
                left = sec.left_margin,
                right = sec.right_margin,
                bottom = sec.bottom_margin
            )
            
            serializer = SectionSerializer(newSec)
            return Response(serializer.data, status=status.HTTP_200_OK)
        except json.JSONDecodeError:
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    return Response(status=status.HTTP_400_BAD_REQUEST)

@api_view(['POST'])
def add_form(request):
    if request.method == 'POST' and request.FILES.get('file'):
        file = request.FILES['file']
        item_name = request.POST.get('itemName')
        try:
            doc = Document(file)
            newForm = DocumentFormat.objects.create(name=item_name )
            for para in doc.paragraphs:  
                paragraph = create_para(para,newForm)
                for run in para.runs: 
                    create_text(run,paragraph)
            serializer = DocumentFormatSerializer(newForm)
            return Response(serializer.data, status=status.HTTP_200_OK)
        except json.JSONDecodeError:
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    return Response(status=status.HTTP_400_BAD_REQUEST)

@api_view(['POST'])
def add_image(request):
    if request.method == 'POST' and request.FILES.get('file'):
        file = request.FILES['file']
        try:
            doc = Document(file)
            for rel in doc.part.rels.values():
                 if "image" in rel.target_ref:
                        img = rel.target_part.blob
                        newImage = Image(name='line',image_data = img)
                        serializer = ImageSerializer(newImage)
                        newImage.save()
            return Response(serializer.data, status=status.HTTP_200_OK)
        except json.JSONDecodeError:
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    return Response(status=status.HTTP_400_BAD_REQUEST)

def create_para(para,doc):
    format = para.paragraph_format
    paragraph = Paragraph.objects.create(
        document_format=doc,
        alignment= para.alignment ,
        tabs=[ [tab.position,tab.alignment,tab.leader] for tab in para.paragraph_format.tab_stops],
        indentation={
            'left': format.left_indent,  # Left indent
            'right': format.right_indent,  # Right indent
            'first_line': format.first_line_indent
        },
        spacing= {
            'before': format.space_before if format.space_before else 0,  # Spacing before
            'after': format.space_after if format.space_after else 0 ,  
            'line': format.line_spacing , # Line spacing
            'line_spacing_rule' : format.line_spacing_rule
        }
    )
    return paragraph
def create_text(run,para):
    Text.objects.create(
            paragraph=para,
            content=run.text,
            bold=run.bold,
            italic=run.italic,
            underline=run.underline,
            font=run.font.name if  run.font.size != None else'TH SarabunIT๙',
            size=run.font.size.pt if run.font.size != None else 16,
            color=run.font.color.rgb 
    )

def create_doc_paragraph(doc, paragraph_data,newline = False):
    if newline:
        paragraph = doc.add_paragraph()
    else:
        paragraph = doc.paragraphs[len(doc.paragraphs)-1]
    paragraph.style = doc.styles['Normal']
    for run_data in paragraph_data.texts.all():
        run = paragraph.add_run(run_data.content)
        if run_data.bold:
            run.bold = True
        if run_data.italic:
            run.italic = True
        if run_data.underline:
            run.underline = run_data.underline
        if run_data.font:
            run.font.name = run_data.font
        if run_data.size:
            run.font.size = Pt(run_data.size)
        '''
        if run_data.color != None:
            run.font.color.rgb =  ImageColor.getrgb(run_data.color)
            '''
            
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = paragraph_data.alignment 
    paragraph_format.left_indent =paragraph_data.indentation['left']
    paragraph_format.right_indent = paragraph_data.indentation['right']
    paragraph_format.first_line_indent = paragraph_data.indentation['first_line']
    paragraph_format.space_before = paragraph_data.spacing['before']
    paragraph_format.space_after = paragraph_data.spacing['after']
    paragraph_format.line_spacing = paragraph_data.spacing['line']
    paragraph_format.line_spacing_rule = paragraph_data.spacing['line_spacing_rule']
    tab_stops = paragraph_format.tab_stops
    tab = paragraph_data.tabs[0] if paragraph_data.tabs else None
    if tab:
        tab_stops.add_tab_stop(tab[0],tab[1],tab[2])
    
    
def create_doc_section(doc, section_data,new):
    if new == True:
        new_sec = doc.add_section()
    else:
        new_sec = doc.sections[0]
    new_sec.left_margin = section_data.left
    new_sec.right_margin = section_data.right
    new_sec.top_margin = section_data.top 
    new_sec.bottom_margin = section_data.bottom

def create_doc_image(doc, pic_data):
    image_stream = BytesIO(pic_data)
    run = doc.add_paragraph().add_run()
    run.add_picture(image_stream)

def create_doc_table(doc,pandas):
    df = pandas
    table = doc.add_table(rows=1, cols=len(df.columns))
    column_scale = [4, 4, 2]
    total_scale = sum(column_scale)
    hdr_cells = table.rows[0].cells
    df.columns = ['รายละเอียดตามข้อกำหนด'	,'รายละเอียดที่พิจารณา'	,'ผลพิจารณา']
    for i, column in enumerate(df.columns):
        column_width_percent = (column_scale[i] / total_scale) * 100
        table.columns[i].width = Inches(column_width_percent / 10)
        hdr_cells[i].text = column
        for para in hdr_cells[i].paragraphs:
                para.style = doc.styles['Normal']
                para.paragraph_format.alignment = 1
                para.paragraph_format.space_before = Pt(6)
                para.paragraph_format.space_after = Pt(6)
                for run in para.runs:
                    run.font.name = 'TH SarabunIT๙'
                    run.font.size = Pt(16)
        hdr_tc = hdr_cells[i]._tc
        hdr_tc.get_or_add_tcPr().append(parse_xml(
                                              r'<w:tcBorders {}>'
                                              r'<w:top w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
                                              r'<w:left w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
                                              r'<w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
                                              r'<w:right w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
                                              r'</w:tcBorders>'.format(nsdecls('w') )))
    for index, row in df.iterrows():   
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)
            for para in row_cells[i].paragraphs:
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
                if i == 2:
                    para.paragraph_format.alignment = 1
                for run in para.runs:
                    run.font.name = 'TH SarabunIT๙'
                    run.font.size = Pt(16)
            data_tc = row_cells[i]._tc
            data_tc.get_or_add_tcPr().append(parse_xml(r'<w:tcBorders {}>'
                                                   r'<w:left w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
                                                   r'<w:right w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
                                                   r'</w:tcBorders>'.format(nsdecls('w'))))
    for cell in table.rows[-1].cells:
        last_row_tc = cell._tc
        last_row_tc.get_or_add_tcPr().append(parse_xml(r'<w:tcBorders {}>'
                                                    r'<w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
                                                    r'</w:tcBorders>'.format(nsdecls('w'))))

def split_dataframe(df):
    result_dfs = []  # List to hold the split DataFrames
    temp_df = []  # Temporary list to accumulate rows for a sub-DataFrame
    line = 0  # Initialize line counter
    for i in range(len(df)):
        line += len(df['consideration'].iloc[i]) / 30
        if i > 0:
            line += 1
        temp_df.append(df.iloc[i])
        
        if line > (20 if not result_dfs else 32):
            result_dfs.append(pd.DataFrame(temp_df))
            temp_df = []  # Reset the temporary DataFrame
            line = 0  # Reset the line counter

        # If it's the last element, check if temp_df needs further splitting
        if i == len(df) - 1 and temp_df:
            last_line = 0
            last_temp = []

            for row in temp_df:
                last_line += len(row['consideration']) / 30
                if last_temp:
                    last_line += 1
                last_temp.append(row)

                if last_line > 21:
                    result_dfs.append(pd.DataFrame(last_temp[:-1]))  # First part
                    result_dfs.append(pd.DataFrame([last_temp[-1]] + temp_df[len(last_temp):]))  # Remaining part
                    break
            else:
                result_dfs.append(pd.DataFrame(temp_df))  # If no split needed

    return result_dfs
def print_page_for_pruduct_result(json_data):
    doc = Document()
    pandas =  pd.DataFrame(json_data['table']) #
    dfs = split_dataframe(pandas)
    for page_num in range(len(dfs)):
        create_doc_section(doc, Section.objects.first(),False if page_num == 0 else True)
        head = DocumentFormat.objects.filter(name = 'product_result_head').first()
        for paragraph_data in head.paragraphs.all():
                    create_doc_paragraph(doc, paragraph_data,True)
        if page_num == 0:
            fristPara = DocumentFormat.objects.filter(name = 'product_result_fristPara').first()
            for paragraph_data in fristPara.paragraphs.all():
                create_doc_paragraph(doc, paragraph_data,True)
        doc.add_paragraph()
        create_doc_table(doc,dfs[page_num])
    tail = DocumentFormat.objects.filter(name = 'product_result_tail').first()
    for paragraph_data in tail.paragraphs.all():
        create_doc_paragraph(doc, paragraph_data,True)
    return doc

def print_product_report():
    doc = Document()
    new_line = False
    create_doc_section(doc, Section.objects.first(),False )
    documentFormat = DocumentFormat.objects.filter(name = 'product').first()
    create_doc_image(doc, Image.objects.filter(name = 'garuda').first().image_data)
    for paragraph_data in documentFormat.paragraphs.all():
        create_doc_paragraph(doc, paragraph_data,new_line)
        new_line = True
    return doc

def print_shop_draw_result():
    doc = print_product_report()
    for para in doc.paragraphs:
        for run in para.runs:
                if 'product' in run.text:
                    run.text = run.text.replace('product', 'sDraw')
    return doc

def print_shop_draw_result():
    doc = Document()
    create_doc_section(doc, Section.objects.first(),False )
    documentFormat = DocumentFormat.objects.filter(name = 'ShopDrawResult').first()
    for paragraph_data in documentFormat.paragraphs.all():
        create_doc_paragraph(doc, paragraph_data,True)
    return doc

def get_doc(json_data):
    if json_data['type'] == 'ผลิตภัณฑ์':
        result = print_page_for_pruduct_result(json_data)
        result_page = fillValue(result,json_data,'result')
        report = print_product_report()
        fillValue(report,json_data,'report',result_page = result_page)  
    if json_data['type'] == 'SHOP DRAWING':
        result = print_shop_draw_result()
        fillValue(result,json_data,'report')
        report = print_shop_draw_result()
        fillValue(result,json_data,'report')
    return [report,result]

def convert_numbers_to_thai(string):
    arabic_to_thai = {
        '0': '๐',  # Thai digit zero
        '1': '๑',  # Thai digit one
        '2': '๒',  # Thai digit two
        '3': '๓',  # Thai digit three อ  
        '4': '๔',  # Thai digit four
        '5': '๕',  # Thai digit five
        '6': '๖',  # Thai digit six
        '7': '๗',  # Thai digit seven
        '8': '๘',  # Thai digit eight
        '9': '๙'   # Thai digit nine
    }
    result = ''.join(arabic_to_thai[char] if char in arabic_to_thai else char for char in string)
    return result

def fillValue(doc,json_data,type,result_page = None):
    ser_type = json_data['type']
    json_copy = dict(json_data)
    for date in ['create_date','contractDate','document_date']:
        y,m,d = json_data[date].split('-')
        day = convert_numbers_to_thai(d)
        month = get_month_in_thai(int(m))
        year = convert_numbers_to_thai(y)
        if date == 'create_date':
            json_copy[date] = f"{month} {year}"
        else:
            json_copy[date] = f"{day} {month} {year}"

    fields = list(SL_TYPEMAP[ser_type ].Meta.fields)
    fields.append('pageN')
    fields = sorted(fields, key=len, reverse=True)
    n_page = len(doc.sections)
    c_page = 1
    for para in doc.paragraphs:
        for run in para.runs:
            for value in fields:
                if value in run.text:
                    if value == 'pageN':
                        if type == 'result':
                            run.text = run.text.replace(value, convert_numbers_to_thai(f"{c_page}/{n_page}"))
                            c_page = c_page +1
                        if type == 'report':
                            run.text = run.text.replace(value, convert_numbers_to_thai(f"{result_page}"))
                        
                    else:
                        if(value == 'departmentCode'):
                            run.text = run.text.replace(value, json_copy[value])
                        else:
                            run.text = run.text.replace(value, convert_numbers_to_thai(json_copy[value]))
    if type == 'result':
        return n_page
    
def get_month_in_thai(month):
    thai_months = [
        "มกราคม", "กุมภาพันธ์", "มีนาคม", "เมษายน", "พฤษภาคม", "มิถุนายน",
        "กรกฎาคม", "สิงหาคม", "กันยายน", "ตุลาคม", "พฤศจิกายน", "ธันวาคม"
    ]
    
    if 1 <= month <= 12:
        return thai_months[month - 1]
    
def convert_docx_to_pdf(docx_buffer):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_docx:
        temp_docx.write(docx_buffer.read())  # Read the bytes from the BytesIO object
        temp_docx_path = temp_docx.name  # Save the file path
    
    output_dir = 'temp'
    os.makedirs(output_dir, exist_ok=True)
    temp_pdf_path = os.path.join(output_dir, os.path.basename(temp_docx_path).replace(".docx", ".pdf"))

    pythoncom.CoInitialize()
    try:
        convert(temp_docx_path, temp_pdf_path)
    finally:
        pythoncom.CoUninitialize()

    with open(temp_pdf_path, 'rb') as pdf_file:
        pdf_buffer = BytesIO(pdf_file.read())
    os.remove(temp_docx_path)
    os.remove(temp_pdf_path)
    

    return pdf_buffer


def get_dum_pd():
    columns = ['รายละเอียดตามข้อกำหนด', 'consideration', 'ผลพิจารณา']
    random_sentences = [
        "This is a test sentence.",
        "The quick brown fox jumps over the lazy dog.",
        "A random sentence for testing.",
        "Python programming is versatile.",
        "Machine learning is fascinating.",
        "Artificial intelligence impacts our lives.",
        "Exploring data science is rewarding.",
        "Natural language processing is useful.",
        "Generating synthetic data is creative.",
        "Structured data makes analysis easier.",
        "Data visualization reveals patterns.",
        "Deep learning models are powerful.",
        "Neural networks simulate human brain functions.",
        "Random text generation can be fun.",
        "Big data requires efficient tools.",
        "Statistical methods support predictions.",
        "Data cleaning is crucial in preprocessing.",
        "Cluster analysis groups similar data.",
        "Dimensionality reduction simplifies data.",
        "Feature selection enhances model performance.",
        "This is a test sentence.",
        "The quick brown fox jumps over the lazy dog.",
        "A random sentence for testing.",
        "Python programming is versatile.",
        "Machine learning is fascinating.",
        "Artificial intelligence impacts our lives.",
        "Exploring data science is rewarding.",
        "Natural language processing is useful.",
        "Generating synthetic data is creative.",
        "Structured data makes analysis easier.",
        "Data visualization reveals patterns.",
        "Deep learning models are powerful.",
        "Neural networks simulate human brain functions.",
        "Random text generation can be fun.",
        "Big data requires efficient tools.",
        "Statistical methods support predictions.",
        "Data cleaning is crucial in preprocessing.",
        "Cluster analysis groups similar data.",
        "Dimensionality reduction simplifies data.",
    ]


    data = {
        'รายละเอียดตามข้อกำหนด': list(range(1, len(random_sentences)+1)),
        'consideration': random_sentences,
        'ผลพิจารณา': list(range(1, len(random_sentences)+1))
    }

    return pd.DataFrame(data, columns=columns)