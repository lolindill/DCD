import json
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor 
import base64
from io import BytesIO
from docx.enum.text import WD_BREAK
from .models import DocumentFormat, Paragraph, Text, Section

def create_paragraph(doc, paragraph_data,frist_line = False):
    if frist_line:
        paragraph = doc.paragraphs[0]
    else:
        paragraph = doc.add_paragraph()
    for run_data in paragraph_data.Text:
        run = paragraph.add_run(run_data.content)
        if run_data.blod:
            run.bold = True
        if run_data.italic:
            run.italic = True
        if run_data.underline:
            run.underline = run_data.underline
        if run_data.font:
            run.font.name = run_data.font
        if run_data.size:
            run.font.size = Pt(run_data.size)
        if run_data.color != None:
            run.font.color.rgb = RGBColor(*run_data.color)
            
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = paragraph_data.alignment 
    paragraph_format.left_indent =paragraph_data.indentation['left']
    paragraph_format.right_indent = paragraph_data.indentation['right']
    paragraph_format.first_line_indent = paragraph_data.indentation['first_line']
    paragraph_format.space_before = paragraph_data.spacing['before']
    paragraph_format.space_after = paragraph_data.spacing['after']
    paragraph_format.line_spacing = paragraph_data.spacing['line']
    paragraph_format.line_spacing_rule = paragraph_data.spacing['line_spacing_rule']
    tab_stops = paragraph_format.tab_stops
    tab = paragraph_data.tabs[0] if paragraph_data.tabs else None
    if tab:
        tab_stops.add_tab_stop(tab[0],tab[1],tab[2])
    
    
def create_section(doc, section_data):
    new_sec = doc.sections[0]
    new_sec.left_margin = section_data.left
    new_sec.right_margin = section_data.right
    new_sec.top_margin = section_data.top 
    new_sec.bottom_margin = section_data.bottom

def create_image(doc, pic_data):
    img_bytes = base64.b64decode(pic_data["image_data"])
    image_stream = BytesIO(img_bytes)
    run = doc.add_paragraph().add_run()
    run.add_picture(image_stream)
    

def get_doc(section, documentFormat):
    doc = Document()
    first_line = True
    '''
    for pic_data in json_data.get('picture', []):
        if pic_data != None:
            create_image(doc, pic_data)
    '''
    
    create_section(doc, Section.objects.first())
    for paragraph_data in documentFormat.Paragraph:
        create_paragraph(doc, paragraph_data,first_line)
        first_line = False
    
    return doc